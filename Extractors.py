"""
extractors.py
Multi-format document text extraction for report analysis:
PDF, DOCX, TXT, CSV, XLSX, JSON.
"""

import csv
import io
import json

import fitz  # PyMuPDF


def extract_pdf(file_bytes: bytes) -> str:
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    return "\n".join(page.get_text() for page in doc)


def extract_docx(file_bytes: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def extract_txt(file_bytes: bytes) -> str:
    return file_bytes.decode("utf-8", errors="ignore")


def extract_csv(file_bytes: bytes) -> str:
    text = file_bytes.decode("utf-8", errors="ignore")
    reader = csv.reader(io.StringIO(text))
    return "\n".join(" | ".join(row) for row in reader)


def extract_xlsx(file_bytes: bytes) -> str:
    import openpyxl

    wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
    lines = []
    for sheet in wb.worksheets:
        lines.append(f"--- Sheet: {sheet.title} ---")
        for row in sheet.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None]
            if cells:
                lines.append(" | ".join(cells))
    return "\n".join(lines)


def extract_json(file_bytes: bytes) -> str:
    data = json.loads(file_bytes.decode("utf-8", errors="ignore"))
    return json.dumps(data, indent=2)


EXTRACTORS = {
    "pdf": extract_pdf,
    "docx": extract_docx,
    "txt": extract_txt,
    "csv": extract_csv,
    "xlsx": extract_xlsx,
    "json": extract_json,
}


def extract_text(filename: str, file_bytes: bytes) -> str:
    ext = filename.rsplit(".", 1)[-1].lower()
    extractor = EXTRACTORS.get(ext)
    if not extractor:
        raise ValueError(f"Unsupported file type: .{ext}")
    return extractor(file_bytes)
